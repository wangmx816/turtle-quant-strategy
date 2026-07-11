#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成 TASK4 Word 报告（姓名+TASK4.docx）。"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

from src.data_fetch import STOCKS, load_all
from src.signals import TurtleConfig
from src.backtest import run_backtest, run_param_grid
from dashboard.build_dashboard import GridConfig, plot_equity, plot_heatmap, plot_signals

ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
OUT_DIR = ROOT / "output"


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.bold = bold


def add_figure(doc: Document, path: Path, caption: str, interpretation: str) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Cm(15))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, caption, bold=True)
    add_para(doc, interpretation)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--name", default="学员", help="报告姓名前缀")
    args = parser.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    stock_data = load_all()
    cfg = TurtleConfig()
    grid = GridConfig()

    for symbol in STOCKS:
        df = stock_data[symbol]
        result = run_backtest(df, cfg)
        grid_df = run_param_grid(df, grid.entry_periods, grid.exit_periods)
        name = STOCKS[symbol]["name"]
        plot_signals(result, name, OUT_DIR / f"signals_{symbol}.png", cfg)
        plot_equity(result, name, OUT_DIR / f"equity_curve_{symbol}.png")
        if not grid_df.empty:
            plot_heatmap(grid_df, name, OUT_DIR / f"heatmap_sharpe_{symbol}.png")

    sym = "002202"
    result = run_backtest(stock_data[sym], cfg)
    grid_df = run_param_grid(stock_data[sym], grid.entry_periods, grid.exit_periods)
    best = grid_df.loc[grid_df["sharpe_ratio"].idxmax()] if not grid_df.empty else None
    m = result["metrics"]

    doc = Document()
    set_style(doc)
    add_para(doc, "海龟交易法则量化回测报告（TASK4）", bold=True)

    add_para(doc, "一、海龟策略核心思想与优势")
    add_para(
        doc,
        "海龟交易法由理查德·丹尼斯于1983年系统化总结，核心思想是：价格突破N日高点时顺势建仓，"
        "跌破M日低点或触发ATR止损时机械离场。其优势在于规则明确、可回测、趋势友好、风险可通过ATR统一度量，"
        "适合流动性较好的趋势性资产；局限是在震荡市中假突破较多，信号具有一定滞后性。",
    )

    add_para(doc, "二、关键概念解释")
    add_para(
        doc,
        "唐奇安通道：入场通道周期N定义过去N日最高价上轨，收盘价向上突破则买入；"
        "离场通道周期M定义过去M日最低价下轨，收盘价向下跌破则卖出。"
        "ATR（平均真实波幅）衡量日内波动，用于计算动态止损距离（止损价=持仓后最高价-系数×ATR）。"
        "止损优先于通道出场，可有效限制单笔亏损。",
    )

    add_para(doc, "三、数据与实现")
    add_para(
        doc,
        f"本项目使用与 quant-strategy 相同的五只A股标的及东方财富前复权日线数据，"
        f"默认参数 entry=20, exit=10, ATR=20, 止损系数=2.0。"
        f"金风科技默认参数回测结果：年化收益{m['annualized_return']}%，夏普{m['sharpe_ratio']}，"
        f"最大回撤{m['max_drawdown']}%，胜率{m['win_rate']}%，交易{m['trade_count']}笔，止损{m['stop_loss_count']}次。",
    )

    add_figure(
        doc,
        OUT_DIR / f"signals_{sym}.png",
        "图1 金风科技海龟策略 — 价格、唐奇安通道与交易信号",
        "图1显示收盘价（灰线）、入场上轨（红虚线）与离场上轨（绿虚线）。绿色上三角为突破买入点，"
        "红色下三角为止损或通道卖出点。可见策略在趋势段持仓，震荡段频繁交易。",
    )

    add_figure(
        doc,
        OUT_DIR / f"heatmap_sharpe_{sym}.png",
        "图2 金风科技入场/出场通道 Sharpe 比率敏感性分析",
        (
            f"图2热力图展示不同 entry/exit 组合下的夏普比率。"
            f"{'最优组合为 entry=' + str(int(best.entry_period)) + ', exit=' + str(int(best.exit_period)) + '，Sharpe=' + f'{best.sharpe_ratio:.3f}' + '。' if best is not None else ''}"
            "颜色偏绿区域表示风险调整后收益更优，说明参数选择对策略表现影响显著，需针对标的做敏感性分析。"
        ),
    )

    add_figure(
        doc,
        OUT_DIR / f"equity_curve_{sym}.png",
        "图3 金风科技策略净值 vs 买入持有基准",
        "图3对比策略净值（红线）与买入持有（蓝线）。若策略曲线在多数时段高于基准，"
        "说明海龟法则在该阶段具有超额收益；反之则需调整参数或考虑市场处于震荡期。",
    )

    add_para(doc, "四、参数调节与使用心得")
    add_para(
        doc,
        "（1）股票类型：工程机械等周期股趋势性较强，适合标准海龟参数；农业、公用事业类震荡较多，"
        "可缩短出场周期M或收紧止损系数。（2）入场N增大：信号减少、滞后增加，适合慢牛行情；"
        "N减小则交易频繁。（3）出场M减小：更快止盈止损，降低回撤但可能错过趋势延续。"
        "（4）综合建议：先用热力图找标的专属参数，再结合 Playground 在线看板验证最新数据表现。",
    )

    add_para(doc, "五、在线看板")
    add_para(
        doc,
        "交互式 Playground 已部署至 https://wangmx816.github.io/turtle-quant-strategy/ ，"
        "支持选择标的、时段与策略参数，GitHub Actions 工作日自动更新行情数据。",
    )

    out = ROOT / f"{args.name}+TASK4.docx"
    doc.save(out)
    print(f"report={out}")


if __name__ == "__main__":
    main()
